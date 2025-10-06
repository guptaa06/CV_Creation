"""
Document generation service - Generate PDF and DOCX resumes
"""
import logging
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from app.models.schemas import ResumeData
from app.config import settings

logger = logging.getLogger(__name__)


class DocumentGenerator:
    """Generate professional resume documents"""

    def __init__(self):
        self.output_dir = Path(settings.OUTPUT_DIR)
        self.output_dir.mkdir(exist_ok=True)

    def generate_pdf(self, resume_data: ResumeData, filename: str = None) -> str:
        """Generate PDF resume"""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"resume_{timestamp}.pdf"

        filepath = self.output_dir / filename

        doc = SimpleDocTemplate(str(filepath), pagesize=letter,
                                rightMargin=0.75*inch, leftMargin=0.75*inch,
                                topMargin=0.75*inch, bottomMargin=0.75*inch)

        story = []
        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'],
                                      fontSize=18, textColor=colors.HexColor('#2C3E50'),
                                      spaceAfter=6, alignment=1)
        heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading2'],
                                       fontSize=14, textColor=colors.HexColor('#2C3E50'),
                                       spaceAfter=6, spaceBefore=12)
        normal_style = styles['Normal']
        normal_style.fontSize = 10

        # Name
        story.append(Paragraph(resume_data.personal_info.name, title_style))
        story.append(Spacer(1, 0.1*inch))

        # Contact Info
        contact_parts = []
        if resume_data.personal_info.email:
            contact_parts.append(resume_data.personal_info.email)
        if resume_data.personal_info.phone:
            contact_parts.append(resume_data.personal_info.phone)
        if resume_data.personal_info.location:
            contact_parts.append(resume_data.personal_info.location)

        contact_text = " | ".join(contact_parts)
        contact_style = ParagraphStyle('Contact', parent=normal_style, alignment=1)
        story.append(Paragraph(contact_text, contact_style))
        story.append(Spacer(1, 0.1*inch))

        # Links
        links = []
        if resume_data.personal_info.linkedin:
            links.append(f"LinkedIn: {resume_data.personal_info.linkedin}")
        if resume_data.personal_info.github:
            links.append(f"GitHub: {resume_data.personal_info.github}")
        if links:
            links_text = " | ".join(links)
            story.append(Paragraph(links_text, contact_style))
            story.append(Spacer(1, 0.15*inch))

        # Summary
        if resume_data.summary:
            story.append(Paragraph("PROFESSIONAL SUMMARY", heading_style))
            story.append(Paragraph(resume_data.summary, normal_style))
            story.append(Spacer(1, 0.1*inch))

        # Skills
        if resume_data.skills:
            story.append(Paragraph("SKILLS", heading_style))
            skills_text = ", ".join(resume_data.skills)
            story.append(Paragraph(skills_text, normal_style))
            story.append(Spacer(1, 0.1*inch))

        # Experience
        if resume_data.experience:
            story.append(Paragraph("PROFESSIONAL EXPERIENCE", heading_style))
            for exp in resume_data.experience:
                job_title = f"<b>{exp.position}</b> | {exp.company}"
                if exp.location:
                    job_title += f" | {exp.location}"
                story.append(Paragraph(job_title, normal_style))

                dates = f"{exp.start_date or ''} - {exp.end_date or 'Present'}"
                story.append(Paragraph(f"<i>{dates}</i>", normal_style))

                for resp in exp.responsibilities:
                    story.append(Paragraph(f"• {resp}", normal_style))

                # Add achievements if present
                if exp.achievements:
                    for achievement in exp.achievements:
                        story.append(Paragraph(f"• {achievement}", normal_style))

                story.append(Spacer(1, 0.1*inch))

        # Education
        if resume_data.education:
            story.append(Paragraph("EDUCATION", heading_style))
            for edu in resume_data.education:
                edu_text = f"<b>{edu.degree}</b>"
                if edu.field_of_study:
                    edu_text += f" in {edu.field_of_study}"
                story.append(Paragraph(edu_text, normal_style))
                story.append(Paragraph(edu.institution, normal_style))
                if edu.gpa:
                    story.append(Paragraph(f"GPA: {edu.gpa}", normal_style))

                # Add academic achievements if present
                if edu.achievements:
                    for achievement in edu.achievements:
                        story.append(Paragraph(f"• {achievement}", normal_style))

                story.append(Spacer(1, 0.1*inch))

        # Projects
        if resume_data.projects:
            story.append(Paragraph("PROJECTS", heading_style))
            for proj in resume_data.projects:
                story.append(Paragraph(f"<b>{proj.name}</b>", normal_style))
                story.append(Paragraph(proj.description, normal_style))
                if proj.technologies:
                    tech_text = f"Technologies: {', '.join(proj.technologies)}"
                    story.append(Paragraph(tech_text, normal_style))
                story.append(Spacer(1, 0.1*inch))

        doc.build(story)
        logger.info(f"Generated PDF resume: {filepath}")
        return str(filepath)

    def generate_docx(self, resume_data: ResumeData, filename: str = None) -> str:
        """Generate DOCX resume"""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"resume_{timestamp}.docx"

        filepath = self.output_dir / filename

        doc = Document()

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Name
        name = doc.add_paragraph(resume_data.personal_info.name)
        name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name.runs[0].font.size = Pt(20)
        name.runs[0].font.bold = True
        name.runs[0].font.color.rgb = RGBColor(44, 62, 80)

        # Contact Info
        contact_parts = []
        if resume_data.personal_info.email:
            contact_parts.append(resume_data.personal_info.email)
        if resume_data.personal_info.phone:
            contact_parts.append(resume_data.personal_info.phone)
        if resume_data.personal_info.location:
            contact_parts.append(resume_data.personal_info.location)

        contact = doc.add_paragraph(" | ".join(contact_parts))
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact.runs[0].font.size = Pt(10)

        # Links
        if resume_data.personal_info.linkedin or resume_data.personal_info.github:
            links_text = " | ".join(filter(None, [
                resume_data.personal_info.linkedin,
                resume_data.personal_info.github
            ]))
            links = doc.add_paragraph(links_text)
            links.alignment = WD_ALIGN_PARAGRAPH.CENTER
            links.runs[0].font.size = Pt(10)

        doc.add_paragraph()  # Spacer

        # Summary
        if resume_data.summary:
            heading = doc.add_heading("PROFESSIONAL SUMMARY", level=2)
            heading.runs[0].font.color.rgb = RGBColor(44, 62, 80)
            doc.add_paragraph(resume_data.summary)

        # Skills
        if resume_data.skills:
            heading = doc.add_heading("SKILLS", level=2)
            heading.runs[0].font.color.rgb = RGBColor(44, 62, 80)
            doc.add_paragraph(", ".join(resume_data.skills))

        # Experience
        if resume_data.experience:
            heading = doc.add_heading("PROFESSIONAL EXPERIENCE", level=2)
            heading.runs[0].font.color.rgb = RGBColor(44, 62, 80)

            for exp in resume_data.experience:
                # Position and company
                exp_para = doc.add_paragraph()
                exp_para.add_run(f"{exp.position}").bold = True
                exp_para.add_run(f" | {exp.company}")

                # Dates
                dates = doc.add_paragraph(f"{exp.start_date or ''} - {exp.end_date or 'Present'}")
                dates.runs[0].italic = True

                # Responsibilities
                for resp in exp.responsibilities:
                    doc.add_paragraph(resp, style='List Bullet')

                doc.add_paragraph()  # Spacer

        # Education
        if resume_data.education:
            heading = doc.add_heading("EDUCATION", level=2)
            heading.runs[0].font.color.rgb = RGBColor(44, 62, 80)

            for edu in resume_data.education:
                edu_para = doc.add_paragraph()
                degree_text = edu.degree
                if edu.field_of_study:
                    degree_text += f" in {edu.field_of_study}"
                edu_para.add_run(degree_text).bold = True

                doc.add_paragraph(edu.institution)
                if edu.gpa:
                    doc.add_paragraph(f"GPA: {edu.gpa}")

        # Projects
        if resume_data.projects:
            heading = doc.add_heading("PROJECTS", level=2)
            heading.runs[0].font.color.rgb = RGBColor(44, 62, 80)

            for proj in resume_data.projects:
                proj_para = doc.add_paragraph()
                proj_para.add_run(proj.name).bold = True
                doc.add_paragraph(proj.description)
                if proj.technologies:
                    doc.add_paragraph(f"Technologies: {', '.join(proj.technologies)}")

        doc.save(str(filepath))
        logger.info(f"Generated DOCX resume: {filepath}")
        return str(filepath)

    def generate(self, resume_data: ResumeData, format: str = "pdf") -> str:
        """Generate resume in specified format"""
        if format.lower() == "pdf":
            return self.generate_pdf(resume_data)
        elif format.lower() in ["docx", "doc"]:
            return self.generate_docx(resume_data)
        else:
            raise ValueError(f"Unsupported format: {format}")


# Global instance
document_generator = DocumentGenerator()
