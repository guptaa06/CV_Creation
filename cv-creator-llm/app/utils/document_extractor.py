"""
Document extraction utilities for PDF, DOCX, and TXT files
"""
import pdfplumber
from docx import Document
from typing import Dict, Any, Optional
import re
import logging

logger = logging.getLogger(__name__)


class DocumentExtractor:
    """Extract text from various document formats"""

    @staticmethod
    def extract_from_txt(file_path: str) -> str:
        """
        Extract text from TXT file

        Args:
            file_path: Path to TXT file

        Returns:
            Extracted text content
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    logger.info(f"Successfully extracted {len(text)} characters from TXT using {encoding} encoding")
                    return text.strip()
                except UnicodeDecodeError:
                    continue

            # If all encodings fail, read as binary and decode with errors ignored
            with open(file_path, 'rb') as f:
                text = f.read().decode('utf-8', errors='ignore')
            logger.info(f"Successfully extracted {len(text)} characters from TXT (with error handling)")
            return text.strip()

        except Exception as e:
            logger.error(f"Error extracting TXT: {str(e)}")
            raise ValueError(f"Failed to extract TXT: {str(e)}")

    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        """
        Extract text from PDF file using pdfplumber

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text content
        """
        try:
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

            logger.info(f"Successfully extracted {len(text)} characters from PDF")
            return text.strip()

        except Exception as e:
            logger.error(f"Error extracting PDF: {str(e)}")
            raise ValueError(f"Failed to extract PDF: {str(e)}")

    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        """
        Extract text from DOCX file

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text content
        """
        try:
            doc = Document(file_path)
            text = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)

            extracted_text = "\n".join(text)
            logger.info(f"Successfully extracted {len(extracted_text)} characters from DOCX")
            return extracted_text.strip()

        except Exception as e:
            logger.error(f"Error extracting DOCX: {str(e)}")
            raise ValueError(f"Failed to extract DOCX: {str(e)}")

    @classmethod
    def extract_text(cls, file_path: str) -> str:
        """
        Extract text from file based on extension

        Args:
            file_path: Path to document file

        Returns:
            Extracted text content
        """
        file_path_lower = file_path.lower()

        if file_path_lower.endswith('.pdf'):
            return cls.extract_from_pdf(file_path)
        elif file_path_lower.endswith('.docx') or file_path_lower.endswith('.doc'):
            return cls.extract_from_docx(file_path)
        elif file_path_lower.endswith('.txt'):
            return cls.extract_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_path}")

    @staticmethod
    def clean_text(text: str) -> str:
        """
        Clean extracted text

        Args:
            text: Raw extracted text

        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)

        # Remove special characters that might interfere with processing
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)

        return text.strip()

    @classmethod
    def extract_and_clean(cls, file_path: str) -> str:
        """
        Extract and clean text from document

        Args:
            file_path: Path to document file

        Returns:
            Clean extracted text
        """
        raw_text = cls.extract_text(file_path)
        return cls.clean_text(raw_text)


def extract_contact_info(text: str) -> Dict[str, Optional[str]]:
    """
    Extract contact information using regex patterns

    Args:
        text: Resume text

    Returns:
        Dictionary with contact information
    """
    contact_info = {
        "email": None,
        "phone": None,
        "linkedin": None,
        "github": None,
    }

    # Email pattern
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    email_match = re.search(email_pattern, text)
    if email_match:
        contact_info["email"] = email_match.group()

    # Phone pattern (various formats)
    phone_pattern = r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    phone_match = re.search(phone_pattern, text)
    if phone_match:
        contact_info["phone"] = phone_match.group()

    # LinkedIn pattern
    linkedin_pattern = r'linkedin\.com/in/[\w-]+'
    linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
    if linkedin_match:
        contact_info["linkedin"] = linkedin_match.group()

    # GitHub pattern
    github_pattern = r'github\.com/[\w-]+'
    github_match = re.search(github_pattern, text, re.IGNORECASE)
    if github_match:
        contact_info["github"] = github_match.group()

    return contact_info
